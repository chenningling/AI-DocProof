#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档处理模块
负责Word文档的读取、分句和批注
"""

import os
import re
import logging
import datetime
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocProcessor:
    """文档处理类，负责Word文档的读取、分句和批注"""
    
    def __init__(self, file_path=None):
        """
        初始化文档处理器
        
        Args:
            file_path: Word文档路径
        """
        self.file_path = file_path
        self.document = None
        self.sentences = []
        self.sentence_positions = []  # 存储每个句子在文档中的位置信息
        
        # 分句规则 - 优化英文句号的处理
        # 中文句号使用全角句号，英文句号需要更严格的匹配条件
        self.sentence_end_pattern = r'(。？！|\.[\s\n]|\?|!)["\'\'\'"]?(?=$|[\s\n])'
        
        # 排除的特殊情况
        self.exclude_patterns = [
            r'\d+\.\d+',  # 小数点
            r'[A-Za-z]\.',  # 英文缩写
            r'Dr\.',  # 博士
            r'Mr\.',  # 先生
            r'Mrs\.',  # 夫人
            r'Ms\.',  # 女士
            r'etc\.',  # 等等
            r'e\.g\.',  # 例如
            r'i\.e\.',  # 也就是
            r'Fig\.',  # 图
            r'Tab\.',  # 表
            r'Eq\.',  # 方程
            r'Ref\.',  # 参考
            r'Vol\.',  # 卷
            r'No\.',  # 编号
            r'p\.',  # 页
            r'pp\.',  # 页
            r'et al\.',  # 等人
            r'\w+\.\w+',  # 包含点的单词，如域名、文件名等
            r'[A-Z]\.[A-Z]',  # 大写字母缩写，如U.S.
            r'Ph\.D',  # 博士学位
            r'\d+\.',  # 序号，如"1."
        ]
        
    def load_document(self, file_path=None):
        """
        加载Word文档
        
        Args:
            file_path: Word文档路径，如果为None则使用初始化时的路径
            
        Returns:
            bool: 加载是否成功
        """
        if file_path:
            self.file_path = file_path
            
        if not self.file_path:
            logger.error("未指定文档路径")
            return False
            
        try:
            self.document = Document(self.file_path)
            logger.info(f"成功加载文档: {self.file_path}")
            return True
        except Exception as e:
            logger.error(f"加载文档失败: {str(e)}")
            return False
    
    def split_into_sentences(self):
        """
        将文档内容分割成句子
        
        Returns:
            list: 句子列表
        """
        if not self.document:
            logger.error("文档未加载")
            return []
            
        self.sentences = []
        self.sentence_positions = []
        
        # 遍历文档中的段落和表格
        for para_index, paragraph in enumerate(self.document.paragraphs):
            if not paragraph.text.strip():
                continue
                
            # 处理段落中的文本
            text = paragraph.text
            current_pos = 0
            
            while current_pos < len(text):
                # 查找句子结束位置
                match = re.search(self.sentence_end_pattern, text[current_pos:])
                
                if match:
                    end_pos = current_pos + match.end()
                    sentence = text[current_pos:end_pos].strip()
                    
                    # 检查是否是排除的特殊情况
                    is_excluded = False
                    
                    # 如果句子太短且以英文句号结尾，可能是错误分割
                    if len(sentence) < 5 and sentence.endswith('.'):
                        is_excluded = True
                    
                    # 检查排除模式
                    for pattern in self.exclude_patterns:
                        if re.search(pattern, sentence):
                            is_excluded = True
                            break
                                
                    if not is_excluded and sentence:
                        self.sentences.append(sentence)
                        # 存储句子位置信息
                        self.sentence_positions.append({
                            'type': 'paragraph',
                            'paragraph_index': para_index,
                            'start': current_pos,
                            'end': end_pos,
                            'text': sentence
                        })
                        
                    current_pos = end_pos
                else:
                    # 如果没有找到句子结束标记，将剩余文本作为一个句子
                    sentence = text[current_pos:].strip()
                    if sentence:
                        self.sentences.append(sentence)
                        self.sentence_positions.append({
                            'type': 'paragraph',
                            'paragraph_index': para_index,
                            'start': current_pos,
                            'end': len(text),
                            'text': sentence
                        })
                    break
        
        # 处理表格中的文本
        for table_index, table in enumerate(self.document.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for para_index, paragraph in enumerate(cell.paragraphs):
                        if not paragraph.text.strip():
                            continue
                            
                        text = paragraph.text
                        current_pos = 0
                        
                        while current_pos < len(text):
                            match = re.search(self.sentence_end_pattern, text[current_pos:])
                            
                            if match:
                                end_pos = current_pos + match.end()
                                sentence = text[current_pos:end_pos].strip()
                                
                                is_excluded = False
                                for pattern in self.exclude_patterns:
                                    if re.search(pattern, sentence):
                                        is_excluded = True
                                        break
                                        
                                if not is_excluded and sentence:
                                    self.sentences.append(sentence)
                                    self.sentence_positions.append({
                                        'type': 'table',
                                        'table_index': table_index,
                                        'row_index': row_index,
                                        'cell_index': cell_index,
                                        'paragraph_index': para_index,
                                        'start': current_pos,
                                        'end': end_pos,
                                        'text': sentence
                                    })
                                    
                                current_pos = end_pos
                            else:
                                sentence = text[current_pos:].strip()
                                if sentence:
                                    self.sentences.append(sentence)
                                    self.sentence_positions.append({
                                        'type': 'table',
                                        'table_index': table_index,
                                        'row_index': row_index,
                                        'cell_index': cell_index,
                                        'paragraph_index': para_index,
                                        'start': current_pos,
                                        'end': len(text),
                                        'text': sentence
                                    })
                                break
        
        logger.info(f"文档分句完成，共 {len(self.sentences)} 个句子")
        return self.sentences
    
    def add_comment(self, sentence_index, comment_text):
        """
        为指定句子添加批注
        
        Args:
            sentence_index: 句子索引
            comment_text: 批注内容
            
        Returns:
            bool: 添加批注是否成功
        """
        if not self.document or sentence_index >= len(self.sentence_positions):
            logger.error("文档未加载或句子索引无效")
            return False
            
        try:
            position = self.sentence_positions[sentence_index]
            
            if position['type'] == 'paragraph':
                # 为段落中的句子添加批注
                paragraph = self.document.paragraphs[position['paragraph_index']]
                self._add_comment_to_paragraph(paragraph, position['start'], position['end'], comment_text)
                
            elif position['type'] == 'table':
                # 为表格中的句子添加批注
                table = self.document.tables[position['table_index']]
                cell = table.rows[position['row_index']].cells[position['cell_index']]
                paragraph = cell.paragraphs[position['paragraph_index']]
                self._add_comment_to_paragraph(paragraph, position['start'], position['end'], comment_text)
                
            logger.info(f"成功为句子添加批注: {self.sentences[sentence_index]}")
            return True
            
        except Exception as e:
            logger.error(f"添加批注失败: {str(e)}")
            return False
    
    def _add_comment_to_paragraph(self, paragraph, start, end, comment_text):
        """
        为段落中的特定文本添加批注
        
        Args:
            paragraph: 段落对象
            start: 文本开始位置
            end: 文本结束位置
            comment_text: 批注内容
        """
        try:
            # 保存原始文本
            text = paragraph.text
            
            # 清除原段落中的所有运行
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            
            # 添加原文内容
            paragraph.add_run(text)
            
            # 添加空格分隔符
            paragraph.add_run(' ')
            
            # 添加高亮显示的批注内容
            comment_run = paragraph.add_run(f"[校对建议: {comment_text}]")
            comment_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
            comment_run.font.bold = True  # 加粗
            comment_run.font.italic = True  # 斜体
            
            logger.info(f"成功为文本添加批注: {text}")
            return True
        except Exception as e:
            logger.error(f"添加批注到段落时出错: {str(e)}")
            return False
    
    # 注意: 我们不再使用此方法，因为我们现在直接在原文后添加高亮的修改建议
    def _create_comment(self, comment_id, comment_text):
        """
        创建批注 (已弃用)
        
        Args:
            comment_id: 批注ID
            comment_text: 批注内容
            
        Returns:
            comment: 批注对象
        """
        # 此方法已不再使用
        return None
    
    def save_document(self, output_path=None):
        """
        保存文档
        
        Args:
            output_path: 输出文档路径，如果为None则自动生成
            
        Returns:
            tuple: (是否成功, 保存路径或错误信息)
        """
        if not self.document:
            logger.error("文档未加载")
            return False, "文档未加载"
            
        try:
            if output_path is None:
                # 自动生成输出路径
                file_dir = os.path.dirname(self.file_path)
                file_name = os.path.basename(self.file_path)
                name, ext = os.path.splitext(file_name)
                
                # 检查是否存在同名文件
                output_path = os.path.join(file_dir, f"{name}_AI校对{ext}")
                if os.path.exists(output_path):
                    # 如果存在同名文件，添加时间戳
                    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
                    output_path = os.path.join(file_dir, f"{name}_AI校对_{timestamp}{ext}")
                    
            # 保存文档
            self.document.save(output_path)
            logger.info(f"成功保存文档: {output_path}")
            return True, output_path
            
        except Exception as e:
            logger.error(f"保存文档失败: {str(e)}")
            return False, f"保存文档失败: {str(e)}"


# 测试代码
if __name__ == "__main__":
    # 测试文档处理
    doc_processor = DocProcessor("test.docx")
    if doc_processor.load_document():
        sentences = doc_processor.split_into_sentences()
        print(f"共找到 {len(sentences)} 个句子")
        
        # 测试添加批注
        if len(sentences) > 0:
            doc_processor.add_comment(0, "这是一个测试批注")
            
        # 保存文档
        success, path = doc_processor.save_document()
        print(f"保存文档: {success}, 路径: {path}")
